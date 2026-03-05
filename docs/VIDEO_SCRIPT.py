#!/usr/bin/env python3
"""Generate VIDEO_SCRIPT.docx - speaking script for project demo video."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[DEMO: {text}]")
    run.bold = True
    run.italic = True
    return p

def main():
    doc = Document()
    doc.add_heading('AI Voice Assistant Pro - Video Demo Script', 0)
    doc.add_paragraph(
        'Use this script while recording your project walkthrough. '
        'Read the text aloud. Follow [DEMO] cues to show the corresponding action on screen.'
    )
    doc.add_paragraph()

    # --- INTRODUCTION ---
    add_heading(doc, '1. Introduction (30 sec)', level=1)
    add_para(doc, (
        "Hi everyone! Today I'm going to show you my project — AI Voice Assistant Pro. "
        "It's an AI-powered voice assistant with a modern web interface that lets you interact "
        "through voice, text, and even images. Let me walk you through what it does and how it works."
    ))
    doc.add_paragraph()

    # --- WHAT IT DOES ---
    add_heading(doc, '2. What It Does (45 sec)', level=1)
    add_para(doc, (
        "The assistant has three main ways to interact. First, voice control — you can speak "
        "directly to it, and it transcribes your speech, sends it to an AI, and speaks the reply back to you. "
        "Second, text input — you type a message and get intelligent responses powered by OpenAI or other LLMs. "
        "Third, image analysis — you upload an image and the assistant uses computer vision to describe "
        "or answer questions about it. It also has built-in tools like weather, calculator, time, and web search "
        "when the LangGraph agent is enabled. So it's a full-featured AI assistant in your browser."
    ))
    doc.add_paragraph()

    # --- LIVE DEMO ---
    add_heading(doc, '3. Live Demo (2–3 min)', level=1)

    add_para(doc, "Let me show you the interface.", bold=True)
    add_note(doc, "Show the main web page at localhost:5002")
    add_para(doc, (
        "This is the main screen. You have the voice control section at the top with Start and Stop buttons. "
        "Below that are cards for Text Input and Image Analysis. And there are feature cards you can click "
        "to quickly try sample prompts. The chat feed shows your conversation history."
    ))
    doc.add_paragraph()

    add_para(doc, "First, let me try voice.", bold=True)
    add_note(doc, "Click Start Listening")
    add_para(doc, (
        "I'll click Start Listening... and speak a question."
    ))
    add_note(doc, "Speak: 'What time is it?' or 'What is the weather in London?'")
    add_para(doc, (
        "Now I'll click Stop. The audio is sent to the server, transcribed, processed by the AI, "
        "and you get both a text reply and the assistant speaks it back to you using text-to-speech."
    ))
    add_note(doc, "Click Stop, wait for response")
    doc.add_paragraph()

    add_para(doc, "Next, text input.", bold=True)
    add_note(doc, "Click a feature card or type in the Text Input box")
    add_para(doc, (
        "I can type something like 'Explain quantum computing in one paragraph' and click Send. "
        "The AI returns a detailed, context-aware response. The conversation history is maintained, "
        "so it remembers what we discussed."
    ))
    add_note(doc, "Type a prompt and click Send")
    doc.add_paragraph()

    add_para(doc, "Finally, image analysis.", bold=True)
    add_note(doc, "Upload an image in the Image Analysis card")
    add_para(doc, (
        "I'll upload an image and optionally add a prompt like 'Describe this image' or 'What objects do you see?'. "
        "When I click Analyze, the image is preprocessed with OpenCV for better quality, then sent to "
        "OpenAI's vision model — GPT-4o — which analyzes it and returns a description or answer. "
        "This is useful for accessibility, content moderation, or quick image understanding."
    ))
    add_note(doc, "Select image, add prompt, click Analyze")
    doc.add_paragraph()

    # --- TECH STACK ---
    add_heading(doc, '4. Tech Stack (45 sec)', level=1)
    add_para(doc, (
        "On the backend, it's built with Python and Flask. Voice recording happens in the browser using "
        "MediaRecorder — the audio is sent to the API, converted with PyDub and ffmpeg, and transcribed "
        "using Google Speech Recognition. For intelligence, it uses OpenAI's API — and optionally "
        "LangGraph for a ReAct agent with tools. Image analysis uses OpenCV for preprocessing and "
        "OpenAI's vision API. Text-to-speech uses the macOS 'say' command or gTTS as fallback. "
        "The UI is a responsive web app with a glassmorphism design, and it can be deployed with Docker."
    ))
    doc.add_paragraph()

    # --- CLOSING ---
    add_heading(doc, '5. Closing (20 sec)', level=1)
    add_para(doc, (
        "That's AI Voice Assistant Pro — voice, text, and image interaction in one AI-powered assistant. "
        "You can clone the repo from GitHub, set up your API keys, and run it locally or in Docker. "
        "Thanks for watching! If you have questions, feel free to reach out. Bye!"
    ))
    doc.add_paragraph()

    add_heading(doc, 'Quick Reference — Demo Cues', level=2)
    add_para(doc, "• Start Listening → speak → Stop Listening")
    add_para(doc, "• Type in Text Input → Send")
    add_para(doc, "• Click feature cards to prefill prompts")
    add_para(doc, "• Upload image → optional prompt → Analyze")

    out_path = 'docs/VIDEO_SCRIPT.docx'
    doc.save(out_path)
    print(f'Saved {out_path}')

if __name__ == '__main__':
    main()
