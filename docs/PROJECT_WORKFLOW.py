#!/usr/bin/env python3
"""Generate PROJECT_WORKFLOW.docx from this script."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text):
    return doc.add_paragraph(text)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def main():
    doc = Document()
    doc.add_heading('AI Voice Assistant Pro - Project Workflow', 0)
    doc.add_paragraph('Documentation of the application workflow and data flow.')

    add_heading(doc, '1. High-Level Overview', level=1)
    add_para(doc, 'The application has three main layers: Web UI (browser), Flask API (routes), and Backend (Chatbot/NLP).')
    add_para(doc, 'User interactions flow: Web UI → Flask API → VoiceChatbot/NLPEngine → Response → Web UI.')

    add_heading(doc, '2. Startup Flow', level=1)
    add_para(doc, 'main.py is the entry point:')
    add_code(doc, '''main.py
  - Load .env
  - Init Flask (templates: web/templates, static: web/static)
  - Init VoiceChatbot()
        - SpeechRecognizer + Microphone (PyAudio)
        - TTS engines (pyttsx3 -> say -> gTTS)
        - NLPEngine (LangGraph agent, LLM, intents)
  - register_routes(app, chatbot)
  - app.run(host, port)''')

    add_heading(doc, '3. Voice Control Flow', level=1)
    add_para(doc, 'User clicks Start Listening, speaks, then clicks Stop.')
    add_code(doc, '''Browser: MediaRecorder captures mic audio (WebM/Opus)
  -> mediaRecorder.onstop -> processAudioRecording()
  -> POST /api/process-audio (FormData: audio blob)
  -> chatbot.listen_for_speech_from_file(audio_file)
        - PyDub: WebM -> WAV (requires ffmpeg)
        - SpeechRecognition: WAV -> text (Google API)
  -> chatbot.process_nlp(text)
  -> chatbot.speak(response)  [TTS: say or gTTS]
  -> Return {text, response} -> Frontend displays in chat''')

    add_heading(doc, '4. Text Input Flow', level=1)
    add_para(doc, 'User types message and clicks Send or presses Enter.')
    add_code(doc, '''POST /api/process-text { "text": "..." }
  -> chatbot.process_nlp(text, user_id)
  -> NLPEngine.process_input(text)
        - _recognize_intent(text)      [regex patterns]
        - _extract_entities(text)      [locations, times, etc.]
        - _analyze_sentiment(text)
        - _update_context(user_id, ...)
        - _generate_response(intent, entities, sentiment, user_id)
              - If USE_LANGCHAIN_AGENT=true: LangGraph agent (fast path or full)
              - Else if USE_LLM=true: LLMIntegration (OpenAI/Claude/Ollama)
              - Else: Built-in intent handlers
  -> Return response -> Frontend adds to conversation''')

    add_heading(doc, '5. Image Analysis Flow', level=1)
    add_para(doc, 'User selects image, optional prompt, clicks Analyze.')
    add_code(doc, '''POST /api/analyze-image (multipart: image, prompt)
  -> chatbot.nlp_engine.analyze_image(image_bytes, prompt)
  -> preprocess_for_vision(image_bytes)  [OpenCV]
        - Resize if > 1024px
        - Denoise
        - CLAHE contrast
        - Encode as JPEG
  -> LLMIntegration.analyze_image(image_bytes, prompt)
        - Pillow: validate and normalize
        - base64 encode
        - OpenAI Vision API (gpt-4o) with image + prompt
  -> Return response -> Frontend adds to conversation''')

    add_heading(doc, '6. LangGraph Agent Flow', level=1)
    add_para(doc, 'When USE_LANGCHAIN_AGENT=true, the LangGraph ReAct agent handles queries with tools.')
    add_code(doc, '''Fast path (simple queries - no LLM):
  - Time: "what time is it" -> _get_current_time() [cached 30s]
  - Calculator: "2+3" -> _safe_calculator()
  - Weather: "weather in X" -> _get_weather() [cached 5min]

Full path (complex queries):
  - create_react_agent.invoke({messages: [HumanMessage]})
  - LLM decides: tool call or answer
  - If tool call -> ToolNode executes -> back to LLM
  - Loop until no more tool calls
  - Return final AIMessage content''')

    add_heading(doc, '7. Component Summary', level=1)
    add_para(doc, 'main.py - Entry point, Flask app, VoiceChatbot init')
    add_para(doc, 'routes.py - API endpoints (/api/process-text, /api/process-audio, /api/analyze-image, etc.)')
    add_para(doc, 'VoiceChatbot - TTS (speak), STT (listen_for_speech_from_file), delegates to NLPEngine')
    add_para(doc, 'NLPEngine - Intent, entities, sentiment, response generation, image analysis')
    add_para(doc, 'LangChainAgent - LangGraph ReAct agent with tools (weather, calculator, time, web_search)')
    add_para(doc, 'LLMIntegration - OpenAI/Claude/Ollama, Vision API')
    add_para(doc, 'image_preprocessing - OpenCV preprocessing before vision API')
    add_para(doc, 'app.js - UI logic, API calls, MediaRecorder, chat display')

    add_heading(doc, '8. Data Flow Summary', level=1)
    add_para(doc, 'Voice:  Mic -> MediaRecorder -> /api/process-audio -> PyDub+SpeechRecognition -> NLPEngine -> TTS -> Response')
    add_para(doc, 'Text:   Input -> /api/process-text -> NLPEngine -> (LangGraph or LLM or built-in) -> Response')
    add_para(doc, 'Image:  File -> /api/analyze-image -> OpenCV preprocess -> Vision API -> Response')

    out_path = 'docs/PROJECT_WORKFLOW.docx'
    doc.save(out_path)
    print(f'Saved {out_path}')

if __name__ == '__main__':
    main()
